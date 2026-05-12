#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Minimal Markdown -> DOCX / PDF renderer.

Handles the subset of Markdown the LLM filter emits:
    # / ## / ### headings
    bullet lists (- foo, * foo)
    numbered lists (1. foo)
    paragraphs
    inline **bold**, *italic*, `code`

This is intentionally small and pure-Python (python-docx + reportlab).
For richer Markdown (tables, links, images) plug in pypandoc later.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from typing import List, Optional

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer


# ---------------------------------------------------------------------------
# Block parsing
# ---------------------------------------------------------------------------

_H_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_B_RE = re.compile(r"^[-*]\s+(.*)$")
_N_RE = re.compile(r"^(\d+)\.\s+(.*)$")


@dataclass
class Block:
    kind: str  # h1, h2, h3, bullet, numbered, paragraph, blank
    text: str = ""
    level: int = 1
    idx: str = ""


def parse_blocks(md: str) -> List[Block]:
    blocks: List[Block] = []
    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            if blocks and blocks[-1].kind != "blank":
                blocks.append(Block(kind="blank"))
            continue
        m = _H_RE.match(line)
        if m:
            level = min(len(m.group(1)), 6)
            blocks.append(Block(kind=f"h{min(level, 3)}", level=level, text=m.group(2).strip()))
            continue
        m = _B_RE.match(line)
        if m:
            blocks.append(Block(kind="bullet", text=m.group(1).strip()))
            continue
        m = _N_RE.match(line)
        if m:
            blocks.append(Block(kind="numbered", idx=m.group(1), text=m.group(2).strip()))
            continue
        blocks.append(Block(kind="paragraph", text=line.strip()))
    return blocks


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

_INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
_INLINE_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_INLINE_CODE = re.compile(r"`([^`]+)`")


def _docx_runs(paragraph, text: str) -> None:
    """Lay out a Markdown line as python-docx Runs, handling **bold**, *italic*, `code`."""
    # Tokenize the line by inline markers, simplest approach: split on a combined regex
    # and decide per chunk. We'll just walk pieces.
    pieces = _split_inline(text)
    for piece in pieces:
        run = paragraph.add_run(piece.text)
        run.bold = piece.bold
        run.italic = piece.italic
        if piece.code:
            run.font.name = "Consolas"
            run.font.size = Pt(10)


@dataclass
class _Piece:
    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False


def _split_inline(text: str) -> List[_Piece]:
    out: List[_Piece] = []
    # Linear pass: alternate between literal text and the next inline match
    i = 0
    while i < len(text):
        b = _INLINE_BOLD.search(text, i)
        it = _INLINE_ITALIC.search(text, i)
        c = _INLINE_CODE.search(text, i)
        nxt = min((m for m in (b, it, c) if m), key=lambda m: m.start(), default=None)
        if nxt is None:
            out.append(_Piece(text=text[i:]))
            break
        if nxt.start() > i:
            out.append(_Piece(text=text[i:nxt.start()]))
        if nxt is b:
            out.append(_Piece(text=b.group(1), bold=True))
        elif nxt is it:
            out.append(_Piece(text=it.group(1), italic=True))
        else:
            out.append(_Piece(text=c.group(1), code=True))
        i = nxt.end()
    return out


def markdown_to_docx(md: str, out_path: str, title: Optional[str] = None) -> str:
    doc = Document()
    if title:
        doc.add_heading(title, level=0)
    blocks = parse_blocks(md)
    for b in blocks:
        if b.kind == "blank":
            continue
        if b.kind.startswith("h"):
            doc.add_heading(b.text, level=int(b.kind[1:]))
        elif b.kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _docx_runs(p, b.text)
        elif b.kind == "numbered":
            p = doc.add_paragraph(style="List Number")
            _docx_runs(p, b.text)
        elif b.kind == "paragraph":
            p = doc.add_paragraph()
            _docx_runs(p, b.text)
    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
# PDF (reportlab)
# ---------------------------------------------------------------------------

def _pdf_inline(text: str) -> str:
    """Convert Markdown inline emphasis to reportlab's mini-HTML."""
    text = _INLINE_BOLD.sub(r"<b>\1</b>", text)
    text = _INLINE_ITALIC.sub(r"<i>\1</i>", text)
    text = _INLINE_CODE.sub(r'<font face="Courier">\1</font>', text)
    return text


def markdown_to_pdf(md: str, out_path: str, title: Optional[str] = None) -> str:
    styles = getSampleStyleSheet()
    body = styles["BodyText"]
    body.spaceAfter = 6
    h_styles = {
        1: styles["Heading1"],
        2: styles["Heading2"],
        3: styles["Heading3"],
    }
    bullet_style = ParagraphStyle("Bullet", parent=body, leftIndent=18, bulletIndent=6)

    doc = SimpleDocTemplate(
        out_path, pagesize=LETTER,
        topMargin=0.75 * inch, bottomMargin=0.75 * inch,
        leftMargin=0.8 * inch, rightMargin=0.8 * inch,
        title=title or "Restructured transcript",
    )

    flow: List = []
    if title:
        flow.append(Paragraph(_pdf_inline(title), styles["Title"]))
        flow.append(Spacer(1, 0.15 * inch))

    blocks = parse_blocks(md)

    # Group consecutive bullets / numbered items so reportlab uses real lists
    i = 0
    while i < len(blocks):
        b = blocks[i]
        if b.kind == "blank":
            flow.append(Spacer(1, 0.08 * inch))
            i += 1
        elif b.kind.startswith("h"):
            level = min(int(b.kind[1:]), 3)
            flow.append(Paragraph(_pdf_inline(b.text), h_styles[level]))
            i += 1
        elif b.kind == "bullet":
            items = []
            while i < len(blocks) and blocks[i].kind == "bullet":
                items.append(ListItem(Paragraph(_pdf_inline(blocks[i].text), bullet_style), leftIndent=18))
                i += 1
            flow.append(ListFlowable(items, bulletType="bullet", start="bulletchar", leftIndent=18))
        elif b.kind == "numbered":
            items = []
            while i < len(blocks) and blocks[i].kind == "numbered":
                items.append(ListItem(Paragraph(_pdf_inline(blocks[i].text), bullet_style), leftIndent=18))
                i += 1
            flow.append(ListFlowable(items, bulletType="1", leftIndent=18))
        elif b.kind == "paragraph":
            flow.append(Paragraph(_pdf_inline(b.text), body))
            i += 1
        else:
            i += 1

    doc.build(flow)
    return out_path
